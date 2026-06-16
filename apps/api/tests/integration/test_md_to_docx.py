"""Integration coverage for POST /api/md-to-docx and GET /api/formats (#60).

No Chromium needed: the DOCX converter is pure Python, so these run everywhere.
"""
from __future__ import annotations

import io

from docx import Document

SAMPLE_MD = b"""---
title: "Docx Test"
---

# Heading One

A paragraph with **bold** and *italic* and `code`.

## Subheading

- item one
- item two

| col a | col b |
| --- | --- |
| 1 | 2 |
"""

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_md_to_docx_returns_a_valid_docx(client):
    resp = client.post(
        "/api/md-to-docx",
        files={"file": ("doc.md", SAMPLE_MD, "text/markdown")},
    )
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"] == DOCX_MIME
    assert resp.headers["content-disposition"] == 'attachment; filename="doc.docx"'

    body = resp.content
    assert body[:4] == b"PK\x03\x04", "not a zip/docx"

    # Opens in python-docx and carries the converted structure.
    doc = Document(io.BytesIO(body))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Heading One" in text
    assert "Subheading" in text
    assert len(doc.tables) == 1
    assert [c.text for c in doc.tables[0].rows[0].cells] == ["col a", "col b"]


def test_md_to_docx_is_deterministic_across_requests(client):
    # The merge-blocker for this feature: same input, same bytes, every run.
    a = client.post("/api/md-to-docx", files={"file": ("doc.md", SAMPLE_MD, "text/markdown")})
    b = client.post("/api/md-to-docx", files={"file": ("doc.md", SAMPLE_MD, "text/markdown")})
    assert a.status_code == b.status_code == 200
    assert a.content == b.content, "DOCX output is not byte-stable across runs"


def test_md_to_docx_rejects_non_markdown(client):
    resp = client.post(
        "/api/md-to-docx",
        files={"file": ("notes.txt", SAMPLE_MD, "text/plain")},
    )
    assert resp.status_code == 400
    assert resp.json()["error"]["code"] == "wrong_file_type"


def test_shipped_format_endpoints_are_real_routes(client):
    # Guards against drift: every shipped pair's declared endpoint must be a
    # route actually registered on the app, so a rename can't silently lie.
    # Read the paths from the OpenAPI schema, a stable public surface, instead
    # of walking app.routes, which mixes in version-specific sub-router wrappers
    # that carry no .path.
    registered = set(client.app.openapi()["paths"])
    resp = client.get("/api/formats")
    assert resp.status_code == 200
    # Drive the loop off the shipped subset and assert it is non-empty, so the
    # check can't pass vacuously if the registry ever returns no shipped pairs
    # (falsegreen C1: assert inside a loop that might never run).
    shipped = [fmt for fmt in resp.json() if fmt["status"] == "shipped"]
    assert shipped, "expected shipped pairs to verify"  # falsegreen: ignore[C6]
    for fmt in shipped:
        assert fmt["endpoint"] in registered, f"{fmt['slug']} endpoint missing: {fmt['endpoint']}"


def test_get_formats_lists_pairs_with_status_and_endpoints(client):
    resp = client.get("/api/formats")
    assert resp.status_code == 200
    data = resp.json()
    by_slug = {f["slug"]: f for f in data}
    assert {"pdf-to-md", "md-to-pdf", "md-to-docx"} <= set(by_slug)

    docx = by_slug["md-to-docx"]
    assert docx["status"] == "shipped"
    assert docx["endpoint"] == "/api/md-to-docx"
    assert docx["output_mime"] == DOCX_MIME

    # Planned pairs are present as metadata with no endpoint.
    planned = [f for f in data if f["status"] in ("roadmap", "wanted")]
    assert planned, "expected planned pairs in the matrix"
    assert all(f["endpoint"] is None for f in planned)
