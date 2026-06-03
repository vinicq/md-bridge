"""Convert structured Markdown to a Word document (.docx) (#60).

Stack:
  markdown (Python) -> HTML -> walked into python-docx elements -> .docx

Usage:
  python convert.py <input.md> -o <output.docx>

Maps the Markdown constructs the project already detects onto Word elements:
headings, paragraphs with bold/italic/inline-code runs, bullet and numbered
lists, block quotes, fenced code blocks, tables, and horizontal rules. YAML
front matter (if present) sets the document title.

Determinism is a hard requirement: the same input must produce a byte-identical
.docx every run, on any machine. python-docx already writes fixed core-property
timestamps, but it stamps every zip member with the current time; the output is
re-packed with a fixed member date so two runs never differ. No clock, no
randomness, no network. Pure-Python (python-docx), no native dependency.
"""
from __future__ import annotations

import argparse
import io
import sys
import zipfile
from html.parser import HTMLParser
from pathlib import Path

import markdown
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401  (kept for future styling)
from docx.shared import Pt

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

MD_EXTENSIONS = ["extra", "sane_lists", "smarty", "toc"]

# A fixed DOS timestamp (the zip epoch, 1980-01-01) for every member, so the
# archive bytes do not depend on when the conversion ran.
_FIXED_ZIP_DATE = (1980, 1, 1, 0, 0, 0)

_MAX_FRONT_MATTER_BYTES = 64 * 1024  # matches the md-to-pdf cap (#150)

# Block tags that open a new paragraph-ish context.
_HEADINGS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
_INLINE = {"strong", "b", "em", "i", "code", "a", "br", "span", "sup", "sub"}


def split_front_matter(text: str) -> tuple[dict, str]:
    """Return (front_matter, body). Mirrors the md-to-pdf splitter (#150)."""
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    raw = text[3:end]
    if len(raw.encode("utf-8")) > _MAX_FRONT_MATTER_BYTES:
        return {}, text
    body = text[end + 4 :].lstrip("\n")
    try:
        data = yaml.safe_load(raw) or {}
    except yaml.YAMLError:
        return {}, text
    return (data if isinstance(data, dict) else {}), body


class _Node:
    """A minimal block node: a tag plus inline runs and/or child blocks."""

    __slots__ = ("tag", "runs", "children", "attrs")

    def __init__(self, tag: str):
        self.tag = tag
        self.runs: list[tuple[str, set[str]]] = []  # (text, styles)
        self.children: list[_Node] = []
        self.attrs: dict[str, str] = {}


class _DomBuilder(HTMLParser):
    """Build a shallow block tree from the markdown-rendered HTML.

    Only the curated element set is modelled; anything else degrades to its text
    content so no input can crash the walker.
    """

    _BLOCKS = {"p", "ul", "ol", "li", "blockquote", "pre", "hr", "table", "thead", "tbody", "tr", "th", "td", *_HEADINGS}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[_Node] = []
        self._stack: list[_Node] = []
        self._inline_styles: list[str] = []

    def _current(self) -> _Node | None:
        return self._stack[-1] if self._stack else None

    def handle_starttag(self, tag: str, attrs):
        if tag in self._BLOCKS:
            node = _Node(tag)
            node.attrs = {k: v or "" for k, v in attrs}
            parent = self._current()
            if parent is not None:
                parent.children.append(node)
            else:
                self.blocks.append(node)
            self._stack.append(node)
        elif tag in _INLINE:
            if tag in ("strong", "b"):
                self._inline_styles.append("bold")
            elif tag in ("em", "i"):
                self._inline_styles.append("italic")
            elif tag == "code":
                self._inline_styles.append("code")
            # a / br / span / sup / sub carry no run style here

    def handle_endtag(self, tag: str):
        if tag in self._BLOCKS:
            # Pop back to and including this tag (tolerate minor nesting slips).
            for i in range(len(self._stack) - 1, -1, -1):
                if self._stack[i].tag == tag:
                    del self._stack[i:]
                    break
        elif tag in ("strong", "b") and "bold" in self._inline_styles:
            self._inline_styles.remove("bold")
        elif tag in ("em", "i") and "italic" in self._inline_styles:
            self._inline_styles.remove("italic")
        elif tag == "code" and "code" in self._inline_styles:
            self._inline_styles.remove("code")

    def handle_data(self, data: str):
        node = self._current()
        if node is None:
            return
        if node.tag in ("ul", "ol", "table", "thead", "tbody", "tr"):
            # Whitespace between structural tags: ignore.
            if data.strip():
                node.runs.append((data, set(self._inline_styles)))
            return
        node.runs.append((data, set(self._inline_styles)))


def _add_runs(paragraph, runs: list[tuple[str, set[str]]]) -> None:
    for text, styles in runs:
        if not text:
            continue
        run = paragraph.add_run(text)
        if "bold" in styles:
            run.bold = True
        if "italic" in styles:
            run.italic = True
        if "code" in styles:
            run.font.name = "Consolas"


def _list_items(node: _Node) -> list[_Node]:
    return [c for c in node.children if c.tag == "li"]


def _emit_block(doc, node: _Node, *, list_style: str | None = None) -> None:
    tag = node.tag
    if tag in _HEADINGS:
        p = doc.add_heading(level=_HEADINGS[tag])
        p.text = ""
        _add_runs(p, node.runs)
    elif tag == "p":
        p = doc.add_paragraph()
        _add_runs(p, node.runs)
    elif tag in ("ul", "ol"):
        style = "List Bullet" if tag == "ul" else "List Number"
        for li in _list_items(node):
            p = doc.add_paragraph(style=style)
            _add_runs(p, li.runs)
            # Nested lists: emit at the same level (flat fallback).
            for child in li.children:
                if child.tag in ("ul", "ol"):
                    _emit_block(doc, child)
    elif tag == "blockquote":
        for child in node.children:
            if child.tag == "p":
                p = doc.add_paragraph(style="Intense Quote")
                _add_runs(p, child.runs)
        if not node.children and node.runs:
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, node.runs)
    elif tag == "pre":
        # Fenced code: the inner <code> text collected as runs; render monospace.
        text = "".join(t for t, _ in node.runs)
        for child in node.children:
            text += "".join(t for t, _ in child.runs)
        p = doc.add_paragraph()
        run = p.add_run(text.rstrip("\n"))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    elif tag == "table":
        _emit_table(doc, node)
    elif tag == "hr":
        # No native HR; a thin empty paragraph keeps the visual break.
        doc.add_paragraph()


def _emit_table(doc, node: _Node) -> None:
    rows: list[_Node] = []
    for section in node.children:
        if section.tag in ("thead", "tbody"):
            rows.extend(c for c in section.children if c.tag == "tr")
        elif section.tag == "tr":
            rows.append(section)
    if not rows:
        return
    ncols = max((len([c for c in r.children if c.tag in ("th", "td")]) for r in rows), default=0)
    if ncols == 0:
        return
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    for r in rows:
        cells = [c for c in r.children if c.tag in ("th", "td")]
        row_cells = table.add_row().cells
        for idx in range(ncols):
            if idx < len(cells):
                text = "".join(t for t, _ in cells[idx].runs)
                row_cells[idx].text = text


def _build_document(md_text: str) -> Document:
    front, body = split_front_matter(md_text)
    html = markdown.markdown(body, extensions=MD_EXTENSIONS)

    doc = Document()
    title = front.get("title")
    if isinstance(title, str) and title.strip():
        doc.core_properties.title = title.strip()
        doc.add_heading(title.strip(), level=0)

    builder = _DomBuilder()
    builder.feed(html)
    builder.close()
    for node in builder.blocks:
        _emit_block(doc, node)
    return doc


def _deterministic_bytes(doc: Document) -> bytes:
    """Serialize the document to byte-stable .docx output.

    python-docx writes fixed core-property timestamps already; the only varying
    part is the per-member zip date (set to the current time at save). Re-pack
    with a fixed member date so the same input yields identical bytes every run.
    """
    raw = io.BytesIO()
    doc.save(raw)
    src = zipfile.ZipFile(io.BytesIO(raw.getvalue()))
    out = io.BytesIO()
    # compresslevel is pinned (not left to the zlib default) so the bytes do not
    # depend on the platform's zlib build or a future CPython default change.
    # Level 6 is the historical default and matches the committed golden.
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as dst:
        for info in src.infolist():  # preserves member order
            new = zipfile.ZipInfo(info.filename, date_time=_FIXED_ZIP_DATE)
            new.compress_type = info.compress_type
            new.external_attr = info.external_attr
            new.internal_attr = info.internal_attr
            new.create_system = 0  # 0 = MS-DOS, stable across OSes
            dst.writestr(new, src.read(info.filename))
    return out.getvalue()


def convert(md_path: Path, docx_path: Path) -> None:
    """Convert a Markdown file to a deterministic .docx at docx_path."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    doc = _build_document(md_text)
    Path(docx_path).write_bytes(_deterministic_bytes(doc))


def convert_bytes(md_bytes: bytes) -> bytes:
    """Convert Markdown bytes to deterministic .docx bytes (no tempfile needed)."""
    doc = _build_document(md_bytes.decode("utf-8"))
    return _deterministic_bytes(doc)


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX.")
    parser.add_argument("md_path", type=Path)
    parser.add_argument("-o", "--output", type=Path, required=True)
    args = parser.parse_args()
    convert(args.md_path, args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
