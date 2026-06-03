"""Regression for the markdown-to-docx skill (#60).

Determinism is the merge-blocker: the same Markdown must produce a byte-identical
.docx on every run and every machine. python-docx writes fixed core-property
timestamps, and the converter re-packs the archive with a fixed member date, so
the output is stable. This suite asserts that byte-stability directly and pins it
against a committed golden.

The golden is byte-comparable because the .docx is fully deterministic; refresh
it after an intended converter change:
    python -m pytest tests/regression/test_md_to_docx_regression.py --update-golden
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest

GOLDEN = Path(__file__).resolve().parent / "golden" / "md-to-docx-sample.docx"

SAMPLE_MD = """---
title: "Regression Sample"
---

# Heading One

A paragraph with **bold** and *italic* and `inline code`.

## Subheading

- item one
- item two

1. first
2. second

> a block quote

```python
print("hello")
```

| col a | col b |
| --- | --- |
| 1 | 2 |
| 3 | 4 |
"""


def test_md_to_docx_is_byte_deterministic(md_to_docx_mod):
    # Two independent conversions of the same input must be byte-identical.
    a = md_to_docx_mod.convert_bytes(SAMPLE_MD.encode("utf-8"))
    b = md_to_docx_mod.convert_bytes(SAMPLE_MD.encode("utf-8"))
    assert a == b, "DOCX output is not byte-stable across runs"
    assert a[:4] == b"PK\x03\x04"


def test_md_to_docx_matches_golden(md_to_docx_mod, update_golden: bool):
    out = md_to_docx_mod.convert_bytes(SAMPLE_MD.encode("utf-8"))

    if update_golden:
        GOLDEN.parent.mkdir(parents=True, exist_ok=True)
        GOLDEN.write_bytes(out)
        pytest.skip(f"updated golden: {GOLDEN.name}")

    assert GOLDEN.exists(), f"missing golden {GOLDEN.name}; run with --update-golden"
    assert out == GOLDEN.read_bytes(), (
        "DOCX output drifted from the golden. If intended (converter or "
        "python-docx change), regenerate with --update-golden."
    )


def test_golden_opens_and_carries_structure():
    # Guard the golden itself: it must be a real, readable .docx.
    from docx import Document

    if not GOLDEN.exists():
        pytest.skip("golden not generated yet")
    doc = Document(io.BytesIO(GOLDEN.read_bytes()))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Heading One" in text and "Subheading" in text
    assert len(doc.tables) == 1
